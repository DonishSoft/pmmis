#!/usr/bin/env python3
"""Export Reports Module documentation to Word (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ARTIFACTS_DIR = os.path.expanduser(
    "~/.gemini/antigravity/brain/6dcbc87c-3bb8-4347-9236-729e9a7d826d"
)
OUTPUT = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Документация_Модуль_Отчёты.docx"
)

SCREENSHOTS = {
    "payments": os.path.join(ARTIFACTS_DIR, "payments_report_1771843457232.png"),
    "indicators": os.path.join(ARTIFACTS_DIR, "indicators_report_1771843476085.png"),
    "employee_kpi": os.path.join(ARTIFACTS_DIR, "employee_kpi_report_1771843495074.png"),
}

# Colors
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1F, 0x2A, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = RGBColor(0x1A, 0x56, 0xDB)
ROW_ALT_BG = RGBColor(0xF3, 0xF4, 0xF6)


def set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders for a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), val.get('sz', '4'))
            el.set(qn('w:color'), val.get('color', '000000'))
            el.set(qn('w:space'), '0')
            tcBorders.append(el)
    tcPr.append(tcBorders)


def make_table(doc, headers, rows, col_widths=None):
    """Create a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "1A56DB")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3F4F6")

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading_styled(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK
    return heading


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
    return p


def add_note(doc, text):
    """Add a note/info paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"ℹ️  {text}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = GRAY


def add_image_safe(doc, path, width=Inches(6.2)):
    """Add image if it exists."""
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Скриншот: {os.path.basename(path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    # ═══════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("МОДУЛЬ «ОТЧЁТЫ»")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Описание функционала")
    run.font.size = Pt(16)
    run.font.color.rgb = GRAY

    doc.add_paragraph()

    system = doc.add_paragraph()
    system.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = system.add_run("PMMIS — Project Management & Monitoring Information System")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run("Проект WSIP — Water Supply and Sanitation Investment Project")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    for _ in range(6):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Февраль 2026 г.")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "Оглавление", level=1)

    toc_items = [
        "1.  Общее описание модуля",
        "2.  Отчёт по платежам",
        "3.  Отчёт по индикаторам",
        "4.  Отчёт по KPI сотрудников",
        "5.  Управление доступом",
        "6.  Техническая архитектура",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style = doc.styles['Normal']
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 1. ОБЩЕЕ ОПИСАНИЕ
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "1. Общее описание модуля", level=1)

    doc.add_paragraph(
        "Модуль «Отчёты» предоставляет аналитические инструменты для мониторинга "
        "ключевых показателей проекта WSIP. Модуль включает три независимых отчёта, "
        "каждый из которых доступен через боковое меню в разделе «Отчётность»."
    )

    add_heading_styled(doc, "Состав модуля", level=2)

    make_table(doc,
        ["№", "Отчёт", "Назначение", "Путь"],
        [
            ["1", "Отчёт по платежам", "Финансовый анализ контрактов", "/Reports/Payments"],
            ["2", "Отчёт по индикаторам", "Прогресс целевых показателей", "/Reports/Indicators"],
            ["3", "KPI сотрудников", "Эффективность персонала", "/Reports/EmployeeKpi"],
        ],
        col_widths=[1, 4.5, 5, 4.5]
    )

    doc.add_paragraph()
    add_heading_styled(doc, "Общие возможности", level=2)

    add_bullet(doc, " — столбчатые, круговые диаграммы (Chart.js)", "📊 Интерактивные графики")
    add_bullet(doc, " — таблицы с раскрывающимися строками", "📋 Детализированные таблицы")
    add_bullet(doc, " — кнопка печати на каждой странице", "🖨️ Печать")
    add_bullet(doc, " — данные загружаются через AJAX с анимацией скелетона", "⚡ Асинхронная загрузка")
    add_bullet(doc, " — каждый отчёт управляется отдельным разрешением", "🔐 Раздельные права")

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 2. ОТЧЁТ ПО ПЛАТЕЖАМ
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "2. Отчёт по платежам", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Путь: ")
    run.bold = True
    run = p.add_run("/Reports/Payments")
    run.font.color.rgb = BLUE

    doc.add_paragraph(
        "Данный отчёт предоставляет полный финансовый анализ запланированных "
        "и фактических платежей по всем контрактам проекта."
    )

    add_image_safe(doc, SCREENSHOTS["payments"])
    doc.add_paragraph()

    # 2.1
    add_heading_styled(doc, "2.1. Сводные карточки", level=2)

    doc.add_paragraph(
        "В верхней части страницы отображаются 4 ключевых показателя:"
    )

    make_table(doc,
        ["Карточка", "Описание"],
        [
            ["Всего контрактов", "Общее количество контрактов в системе"],
            ["Запланировано", "Суммарная плановая стоимость всех контрактов (USD)"],
            ["Оплачено", "Суммарная стоимость оплаченных платежей (USD) с прогресс-баром"],
            ["Остаток", "Разница между планом и фактической оплатой"],
        ],
        col_widths=[5, 12]
    )

    doc.add_paragraph()

    # 2.2
    add_heading_styled(doc, "2.2. Графики", level=2)

    add_heading_styled(doc, "Столбчатая диаграмма «Планируемые vs Оплаченные»", level=3)
    add_bullet(doc, "Сгруппированная столбчатая диаграмма (Bar Chart)", "Тип: ")
    add_bullet(doc, "Номера контрактов", "Ось X: ")
    add_bullet(doc, "Суммы в USD", "Ось Y: ")
    add_bullet(doc, "Синий — запланированная сумма, Зелёный — оплаченная сумма", "Данные: ")

    add_heading_styled(doc, "Круговая диаграмма «Распределение оплат»", level=3)
    add_bullet(doc, "Кольцевая диаграмма (Doughnut Chart)", "Тип: ")
    add_bullet(doc, "Доля оплаченных сумм по каждому контракту", "Данные: ")

    # 2.3
    add_heading_styled(doc, "2.3. Детализация по контрактам", level=2)

    doc.add_paragraph("Основная таблица содержит следующие столбцы:")

    make_table(doc,
        ["Столбец", "Описание"],
        [
            ["Контракт", "Номер контракта (ссылка на редактирование) и код проекта"],
            ["Подрядчик", "Наименование подрядчика"],
            ["Валюта", "USD или TJS с обменным курсом при подписании"],
            ["План", "Плановая сумма в TJS и эквивалент в USD для сомони-контрактов"],
            ["Оплачено", "Фактически оплачено (в TJS + эквивалент USD)"],
            ["Курсовая разница", "▲ прибыль / ▼ убыток от изменения курса USD/TJS"],
            ["Остаток", "Неоплаченная разница"],
            ["%", "Процент выполнения с прогресс-баром"],
        ],
        col_widths=[4, 13]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "Раскрывающаяся детализация платежей", level=3)
    doc.add_paragraph(
        "При клике на строку контракта раскрывается вложенная таблица с историей "
        "каждого платежа, содержащая:"
    )

    make_table(doc,
        ["Поле", "Описание"],
        [
            ["Дата", "Дата совершения платежа"],
            ["Тип", "Аванс / Промежуточный / Окончательный / Удержание"],
            ["Статус", "Ожидает / Одобрен / Оплачен / Отклонён"],
            ["Сумма (TJS)", "Сумма в сомони (для TJS-контрактов)"],
            ["Курс оплаты", "Курс НБТ на дату платежа"],
            ["USD (по курсу контракта)", "Сумма в USD по курсу на дату подписания"],
            ["USD (по курсу оплаты)", "Сумма в USD по курсу на дату оплаты"],
            ["Курсовая разница", "Разница между двумя USD-суммами"],
        ],
        col_widths=[5.5, 11.5]
    )

    doc.add_paragraph()

    # 2.4
    add_heading_styled(doc, "2.4. Расчёт курсовой разницы", level=2)

    doc.add_paragraph(
        "Для контрактов в TJS (сомони) система автоматически рассчитывает "
        "курсовую разницу по следующим формулам:"
    )

    formulas = [
        ("USD по курсу контракта", "Сумма TJS ÷ Курс при подписании"),
        ("USD по курсу оплаты", "Сумма TJS ÷ Курс НБТ на дату платежа"),
        ("Курсовая разница", "USD по курсу оплаты − USD по курсу контракта"),
    ]
    make_table(doc,
        ["Показатель", "Формула"],
        formulas,
        col_widths=[6, 11]
    )

    doc.add_paragraph()
    add_bullet(doc, " (зелёный) — курс вырос, экономия в USD", "▲ Положительная")
    add_bullet(doc, " (красный) — курс упал, дополнительные расходы", "▼ Отрицательная")

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 3. ОТЧЁТ ПО ИНДИКАТОРАМ
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "3. Отчёт по индикаторам", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Путь: ")
    run.bold = True
    run = p.add_run("/Reports/Indicators")
    run.font.color.rgb = BLUE

    doc.add_paragraph(
        "Данный отчёт обеспечивает мониторинг прогресса достижения целевых "
        "показателей (индикаторов) проекта WSIP."
    )

    add_image_safe(doc, SCREENSHOTS["indicators"])
    doc.add_paragraph()

    # 3.1
    add_heading_styled(doc, "3.1. Сводные карточки", level=2)

    make_table(doc,
        ["Карточка", "Описание"],
        [
            ["Всего индикаторов", "Общее количество зарегистрированных индикаторов проекта"],
            ["Выполнено", "Количество индикаторов, достигших 100% и более"],
            ["В процессе", "Количество индикаторов с частичным прогрессом (0–99%)"],
            ["Общий прогресс", "Средний процент достижения по всем индикаторам"],
        ],
        col_widths=[5, 12]
    )

    doc.add_paragraph()

    # 3.2
    add_heading_styled(doc, "3.2. График прогресса", level=2)

    add_bullet(doc, "Сгруппированная столбчатая диаграмма (Bar Chart)", "Тип: ")
    add_bullet(doc, "Коды индикаторов (IR-1, IR-1.1, PDO-1 и т.д.)", "Ось X: ")
    add_bullet(doc, "Синий — целевое значение, Зелёный — достигнутое значение", "Данные: ")

    # 3.3
    add_heading_styled(doc, "3.3. Детализация по индикаторам", level=2)

    make_table(doc,
        ["Столбец", "Описание"],
        [
            ["Код", "Уникальный код индикатора (в виде бейджа)"],
            ["Индикатор", "Название индикатора и его категория"],
            ["Ед. изм.", "Единица измерения (количество, Да/Нет, % и т.д.)"],
            ["Цель", "Целевое значение"],
            ["Достигнуто", "Фактически достигнутое значение"],
            ["Остаток", "Разница между целью и достижением"],
            ["Прогресс", "Визуальный прогресс-бар с процентом"],
        ],
        col_widths=[4, 13]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "Цветовая индикация прогресса", level=3)

    make_table(doc,
        ["Прогресс", "Цвет", "Значение"],
        [
            ["≥ 100%", "Зелёный", "Индикатор достигнут"],
            ["50–99%", "Синий", "Активный прогресс"],
            ["1–49%", "Жёлтый", "Начальный этап"],
            ["0%", "Серый", "Не начат"],
        ],
        col_widths=[4, 4, 9]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "Раскрывающаяся детализация по контрактам", level=3)

    doc.add_paragraph(
        "При клике на строку индикатора раскрывается вложенная таблица "
        "с разбивкой вклада каждого контракта:"
    )

    make_table(doc,
        ["Поле", "Описание"],
        [
            ["Контракт", "Номер контракта (ссылка на редактирование)"],
            ["Подрядчик", "Наименование исполнителя"],
            ["Цель", "Целевое значение для данного контракта"],
            ["Достигнуто", "Фактическое значение"],
            ["%", "Процент выполнения (цветной бейдж)"],
        ],
        col_widths=[4, 13]
    )

    doc.add_paragraph()

    # 3.4
    add_heading_styled(doc, "3.4. Источники данных", level=2)

    doc.add_paragraph("Данные по индикаторам формируются из двух источников:")
    add_bullet(doc, " — целевые значения, привязанные к контрактам", "Индикаторы контрактов (ContractIndicator)")
    add_bullet(doc, " — фактические значения, вносимые через отчёты АВР", "Прогресс по АВР (ContractIndicatorProgress)")

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 4. KPI СОТРУДНИКОВ
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "4. Отчёт по KPI сотрудников", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Путь: ")
    run.bold = True
    run = p.add_run("/Reports/EmployeeKpi")
    run.font.color.rgb = BLUE

    doc.add_paragraph(
        "Данный отчёт предоставляет анализ эффективности сотрудников "
        "по назначенным задачам и активности в системе."
    )

    add_image_safe(doc, SCREENSHOTS["employee_kpi"])
    doc.add_paragraph()

    # 4.1
    add_heading_styled(doc, "4.1. Сводные карточки", level=2)

    make_table(doc,
        ["Карточка", "Описание"],
        [
            ["Сотрудников", "Общее количество активных пользователей (с указанием тех, кто имеет задачи)"],
            ["Всего задач", "Общее количество задач в системе (с указанием выполненных)"],
            ["Среднее выполнение", "Средний процент выполнения задач среди сотрудников"],
            ["Просрочено", "Количество просроченных задач, требующих внимания"],
        ],
        col_widths=[5, 12]
    )

    doc.add_paragraph()

    # 4.2
    add_heading_styled(doc, "4.2. Графики", level=2)

    add_heading_styled(doc, "Горизонтальная диаграмма «Задачи по сотрудникам»", level=3)
    add_bullet(doc, "Стековая горизонтальная диаграмма (Stacked Horizontal Bar)", "Тип: ")
    add_bullet(doc, "Имена сотрудников (только с задачами)", "Ось Y: ")
    add_bullet(doc, "Зелёный — выполнено, Синий — в работе, Красный — просрочено, Серый — прочие", "Данные: ")

    add_heading_styled(doc, "Кольцевая диаграмма «Статус задач»", level=3)
    add_bullet(doc, "Doughnut Chart", "Тип: ")
    add_bullet(doc, "Общее распределение задач: Выполнено / В работе / Просрочено / Прочие", "Данные: ")

    # 4.3
    add_heading_styled(doc, "4.3. Детализация по сотрудникам", level=2)

    make_table(doc,
        ["Столбец", "Описание"],
        [
            ["Сотрудник", "Аватар (инициал), ФИО и email"],
            ["Должность", "Должность сотрудника"],
            ["Всего задач", "Общее количество назначенных задач"],
            ["Выполнено", "Количество завершённых задач (зелёный бейдж)"],
            ["В работе", "Количество задач в активной работе (синий бейдж)"],
            ["Просрочено", "Количество просроченных задач (красный бейдж)"],
            ["АВР", "Количество поданных актов выполненных работ"],
            ["Выполнение", "Процент выполнения с прогресс-баром"],
        ],
        col_widths=[4, 13]
    )

    doc.add_paragraph()

    # 4.4
    add_heading_styled(doc, "4.4. Источники данных", level=2)

    make_table(doc,
        ["Метрика", "Источник"],
        [
            ["Задачи", "Таблица ProjectTasks (по полю AssigneeId)"],
            ["Статус задач", "Поле Status (Completed, InProgress и др.)"],
            ["Просроченные", "Поле IsOverdue (вычисляемое)"],
            ["АВР", "Таблица WorkProgresses (по полю SubmittedByUserId)"],
            ["Активные сотрудники", "Таблица AspNetUsers (поле IsActive = true)"],
        ],
        col_widths=[5, 12]
    )

    doc.add_paragraph()

    # 4.5
    add_heading_styled(doc, "4.5. Формула расчёта KPI", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Процент выполнения = (Выполненные задачи ÷ Всего задач) × 100%")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph("Цветовая индикация:")
    add_bullet(doc, " ≥ 80% — высокая эффективность", "🟢")
    add_bullet(doc, " ≥ 50% — удовлетворительно", "🔵")
    add_bullet(doc, " > 0% — требует улучшения", "🟡")
    add_bullet(doc, " 0% — без задач", "⚪")

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 5. УПРАВЛЕНИЕ ДОСТУПОМ
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "5. Управление доступом", level=1)

    doc.add_paragraph(
        "Каждый отчёт управляется отдельным разрешением через систему RoleMenuPermission. "
        "Это позволяет гибко настраивать доступ к каждому типу отчёта для разных ролей."
    )

    add_heading_styled(doc, "Ключи разрешений", level=2)

    make_table(doc,
        ["Отчёт", "Ключ разрешения (MenuKey)", "Описание в интерфейсе"],
        [
            ["Платежи", "ReportsPayments", "Отчёт: Платежи"],
            ["Индикаторы", "ReportsIndicators", "Отчёт: Индикаторы"],
            ["KPI сотрудников", "ReportsEmployeeKpi", "Отчёт: KPI Сотрудников"],
        ],
        col_widths=[4, 5.5, 5.5]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "Настройка прав доступа", level=2)

    doc.add_paragraph(
        "Администратор может настроить доступ к каждому отчёту через: "
        "Меню → Роли → [Выбрать роль] → Разрешения"
    )

    doc.add_paragraph(
        "Для каждого отчёта доступен один тип разрешения:"
    )
    add_bullet(doc, " — доступ к просмотру отчёта", "Просмотр (View)")

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # 6. ТЕХНИЧЕСКАЯ АРХИТЕКТУРА
    # ═══════════════════════════════════════════
    add_heading_styled(doc, "6. Техническая архитектура", level=1)

    add_heading_styled(doc, "Структура файлов", level=2)

    make_table(doc,
        ["Файл", "Назначение", "Строк"],
        [
            ["ReportsController.cs", "Контроллер — бизнес-логика всех отчётов", "258"],
            ["PaymentReportViewModel.cs", "ViewModel для отчёта по платежам", "—"],
            ["IndicatorReportViewModel.cs", "ViewModel для отчёта по индикаторам", "—"],
            ["EmployeeKpiReportViewModel.cs", "ViewModel для KPI сотрудников", "—"],
            ["Payments.cshtml", "Каркас страницы (AJAX loader)", "97"],
            ["_PaymentsData.cshtml", "Partial view — данные и графики платежей", "424"],
            ["Indicators.cshtml", "Каркас страницы (AJAX loader)", "67"],
            ["_IndicatorsData.cshtml", "Partial view — данные и графики индикаторов", "261"],
            ["EmployeeKpi.cshtml", "Каркас страницы (AJAX loader)", "68"],
            ["_EmployeeKpiData.cshtml", "Partial view — данные и графики KPI", "271"],
        ],
        col_widths=[5.5, 8.5, 2]
    )

    doc.add_paragraph()

    add_heading_styled(doc, "Паттерн загрузки данных", level=2)

    doc.add_paragraph("Все три отчёта используют единый паттерн асинхронной загрузки:")

    steps = [
        "Пользователь открывает страницу отчёта",
        "Отображается скелетон-плейсхолдер (анимация загрузки)",
        "JavaScript отправляет AJAX-запрос на серверный endpoint (*Data)",
        "Сервер выполняет запрос к PostgreSQL и формирует Partial View",
        "HTML-ответ вставляется на страницу, скелетон скрывается",
        "Инициализируются интерактивные Chart.js графики",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        run.font.size = Pt(10)

    doc.add_paragraph()

    add_heading_styled(doc, "Используемые библиотеки", level=2)

    make_table(doc,
        ["Библиотека", "Версия", "Назначение"],
        [
            ["Chart.js", "4.4.4", "Интерактивные диаграммы (Bar, Doughnut)"],
            ["Bootstrap 5", "—", "Верстка, таблицы, collapse-компоненты"],
            ["Bootstrap Icons", "—", "Иконки интерфейса"],
        ],
        col_widths=[4, 3, 10]
    )

    # ═══════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Конец документа —")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True

    # Save
    doc.save(OUTPUT)
    print(f"✅ Документ сохранён: {OUTPUT}")


if __name__ == "__main__":
    build_document()
