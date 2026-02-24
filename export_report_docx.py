#!/usr/bin/env python3
"""
PMMIS Report Export to DOCX
Generates a professional Microsoft Word document with tables, embedded images, and formatting.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Configuration ---
BRAIN_DIR = os.path.expanduser(
    "~/.gemini/antigravity/brain/6dcbc87c-3bb8-4347-9236-729e9a7d826d"
)
OUTPUT_PATH = os.path.join(
    os.path.expanduser("~/Documents/Projects/WB Project PMMIS"),
    "ОТЧЁТ_PMMIS.docx"
)

SCREENSHOTS = {
    "login": os.path.join(BRAIN_DIR, "login_page_1771814257973.png"),
    "dashboard": os.path.join(BRAIN_DIR, "dashboard_page_1771814283552.png"),
    "contracts": os.path.join(BRAIN_DIR, "contracts_list_page_1771814322177.png"),
    "payments": os.path.join(BRAIN_DIR, "payments_page_1771814350528.png"),
    "avr": os.path.join(BRAIN_DIR, "avr_page_1771814353065.png"),
    "procurement": os.path.join(BRAIN_DIR, "procurement_page_1771814355478.png"),
    "tasks": os.path.join(BRAIN_DIR, "tasks_page_1771814399203.png"),
    "geography": os.path.join(BRAIN_DIR, "geography_page_1771814401640.png"),
    "projects": os.path.join(BRAIN_DIR, "projects_page_1771814404051.png"),
    "indicators": os.path.join(BRAIN_DIR, "indicators_page_1771814454553.png"),
    "work_progress": os.path.join(BRAIN_DIR, "work_progress_page_1771814456980.png"),
    "contractors": os.path.join(BRAIN_DIR, "contractors_page_1771814459425.png"),
}

# --- Helpers ---

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=10, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1B3A5C"):
    """Create a professional table with styled header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            set_cell_text(cell, cell_text, size=10)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")

    # Set column widths if specified
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    return table


def add_heading(doc, text, level=1):
    """Add heading with custom formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return heading


def add_screenshot(doc, key, caption, width=6.2):
    """Add screenshot with caption."""
    path = SCREENSHOTS.get(key)
    if path and os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Рис. {caption}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()  # spacing
    else:
        doc.add_paragraph(f"[Скриншот: {caption}]")


def add_info_box(doc, text, box_type="info"):
    """Add a highlighted info box."""
    p = doc.add_paragraph()
    if box_type == "important":
        prefix = "⚠️ "
    elif box_type == "note":
        prefix = "📌 "
    else:
        prefix = "ℹ️ "
    run = p.add_run(prefix + text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x55, 0x88)


# --- Main Document Generation ---

def create_report():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ========== TITLE PAGE ==========
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ОТЧЁТ О ПРОДЕЛАННЫХ РАБОТАХ")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Разработка информационной системы управления\n"
        "проектами и мониторинга — PMMIS\n"
        "(Project Management & Monitoring Information System)"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x66, 0x88)

    for _ in range(4):
        doc.add_paragraph()

    # Title page info table
    info_data = [
        ("Проект:", "WSIP — Water Supply and Sanitation Investment Project\n"
                     "(Проект инвестирования в водоснабжение и санитарию)"),
        ("Заказчик:", "Всемирный Банк (World Bank Group) / ЦУП / PMU"),
        ("Исполнитель:", "DonishSoft"),
        ("Дата:", "Февраль 2026"),
    ]

    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_data):
        set_cell_text(info_table.rows[i].cells[0], label, bold=True, size=12)
        set_cell_text(info_table.rows[i].cells[1], value, size=12)
        info_table.rows[i].cells[0].width = Cm(3.5)
        info_table.rows[i].cells[1].width = Cm(12)

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_heading(doc, "Содержание", level=1)

    toc_items = [
        "1. Введение",
        "2. Назначение и цели системы",
        "3. Архитектура системы",
        "4. Технологический стек",
        "5. Модульная структура PMMIS",
        "6. Описание модулей",
        "    6.1. Dashboard (Панель управления)",
        "    6.2. Модуль «Проекты»",
        "    6.3. Модуль «Контракты»",
        "    6.4. Модуль «Подрядчики»",
        "    6.5. Модуль «АВР» (Акты выполненных работ)",
        "    6.6. Модуль «Прогресс работ»",
        "    6.7. Модуль «Платежи»",
        "    6.8. Модуль «План закупок»",
        "    6.9. Модуль «Индикаторы / KPI»",
        "    6.10. Модуль «Задачи»",
        "    6.11. Модуль «Уведомления»",
        "    6.12. Модуль «География»",
        "    6.13. Модуль «Пользователи и Роли»",
        "    6.14. Модуль «Справочники»",
        "7. Взаимодействие модулей",
        "8. Соответствие ТЗ",
        "9. Заключение",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ========== 1. ВВЕДЕНИЕ ==========
    add_heading(doc, "1. Введение", level=1)

    doc.add_paragraph(
        "Настоящий отчёт представляет результаты работ по разработке информационной системы "
        "PMMIS (Project Management & Monitoring Information System) в рамках проекта "
        "WSIP (Water Supply and Sanitation Investment Project), финансируемого Всемирным Банком."
    )
    doc.add_paragraph(
        "Система разработана в соответствии с Финальной Концепцией Системы (обновлённая версия, "
        "ноябрь 2025) и призвана обеспечить прозрачность, дисциплину и контроль при реализации "
        "инвестиционных проектов водоснабжения в Республике Таджикистан."
    )

    add_info_box(doc, 
        "Центральная логика PMMIS строится вокруг ключевой управленческой цепочки:\n"
        "Контракт → АВР → Платёж → Обязательства → Остаток",
        "important"
    )

    doc.add_paragraph(
        "Эта цепочка отражает реальный управленческий процесс ЦУП и PMU и является "
        "основной зоной риска при реализации проектов. PMMIS не является ERP, не заменяет "
        "бухгалтерию и не предназначена для полной автоматизации процессов учреждения."
    )

    # ========== 2. НАЗНАЧЕНИЕ И ЦЕЛИ ==========
    add_heading(doc, "2. Назначение и цели системы", level=1)

    add_heading(doc, "2.1. Основное назначение", level=2)
    doc.add_paragraph(
        "PMMIS — это лёгкая, практичная система управления исполнением контрактов, "
        "АВР, платежами и закупками, с возможностью учёта технических contract-linked "
        "индикаторов и ведением базовой отчетности по WB-индикаторам."
    )

    add_heading(doc, "2.2. Ключевые цели", level=2)

    goals = [
        ("1", "Обеспечение прозрачности цепочки «Контракт → АВР → Платёж → Остаток»", "✅ Реализовано"),
        ("2", "Автоматический расчёт KPI по контрактам, закупкам и платежам", "✅ Реализовано"),
        ("3", "Управление contract-linked индикаторами (автоматический расчёт)", "✅ Реализовано"),
        ("4", "Ведение outcome-level индикаторов WB (ручной ввод)", "✅ Реализовано"),
        ("5", "Мониторинг рисков через «Управленческие сигналы» (красные флаги)", "✅ Реализовано"),
        ("6", "Минимальная нагрузка на пользователей", "✅ Реализовано"),
        ("7", "Многоязычность (русский, таджикский, английский)", "✅ Реализовано"),
        ("8", "Экспорт в Excel/PDF", "✅ Реализовано"),
    ]
    add_styled_table(doc, ["№", "Цель", "Статус"], goals, col_widths=[1, 11, 4])

    # ========== 3. АРХИТЕКТУРА ==========
    doc.add_page_break()
    add_heading(doc, "3. Архитектура системы", level=1)

    add_heading(doc, "3.1. Архитектурный подход — Clean Architecture", level=2)
    doc.add_paragraph(
        "Система построена на принципах Clean Architecture (Чистая Архитектура), "
        "обеспечивающих разделение ответственности, тестируемость и масштабируемость."
    )

    # Architecture layers table
    arch_layers = [
        ("PMMIS.Web\n(Presentation)", "Пользовательский интерфейс", "MVC Controllers, Razor Views,\nJavaScript, Syncfusion EJ2"),
        ("PMMIS.Application\n(Business Logic)", "Бизнес-логика и интерфейсы сервисов", "Interfaces, DTOs,\nService Contracts"),
        ("PMMIS.Infrastructure\n(Data Access)", "Инфраструктура доступа к данным", "EF Core DbContext (Npgsql),\nMigrations, ASP.NET Identity"),
        ("PMMIS.Domain\n(Core)", "Ядро системы — сущности\nи бизнес-правила", "28 сущностей: Contract, Payment,\nProcurementPlan, ProjectTask,\nIndicator, Geography и др."),
    ]
    add_styled_table(doc, ["Слой", "Назначение", "Ключевые компоненты"], arch_layers, col_widths=[4.5, 5, 7])

    add_heading(doc, "3.2. Базовые паттерны доменной модели", level=2)
    doc.add_paragraph(
        "Все сущности системы наследуют от базовых классов, обеспечивая единообразие:"
    )
    p = doc.add_paragraph()
    run = p.add_run("BaseEntity")
    run.bold = True
    p.add_run(" — минимальное ядро с полями Id, CreatedAt, UpdatedAt")

    p = doc.add_paragraph()
    run = p.add_run("LocalizedEntity")
    run.bold = True
    p.add_run(" — поддержка трёхъязычности (NameRu, NameTj, NameEn) с утилитарным методом GetName(lang)")

    # ========== 4. ТЕХНОЛОГИЧЕСКИЙ СТЕК ==========
    add_heading(doc, "4. Технологический стек", level=1)

    add_heading(doc, "4.1. Серверная часть (Backend)", level=2)
    backend = [
        (".NET", "10.0", "Платформа разработки"),
        ("ASP.NET Core MVC", "10.0", "Веб-фреймворк"),
        ("Entity Framework Core", "10.0", "ORM для работы с БД"),
        ("Npgsql", "—", "Провайдер PostgreSQL для EF Core"),
        ("ASP.NET Core Identity", "—", "Аутентификация и авторизация"),
        ("Serilog", "—", "Структурированное логирование"),
    ]
    add_styled_table(doc, ["Технология", "Версия", "Назначение"], backend, col_widths=[5, 2.5, 8.5])

    add_heading(doc, "4.2. Клиентская часть (Frontend)", level=2)
    frontend = [
        ("Razor Views (CSHTML)", "Шаблонизатор представлений"),
        ("Bootstrap 5", "CSS-фреймворк для адаптивной вёрстки"),
        ("Syncfusion Essential Studio (EJ2)", "Enterprise-grade UI компоненты (DataGrid, Scheduler, Kanban и др.)"),
        ("JavaScript (ES6+)", "Интерактивная логика на клиенте"),
    ]
    add_styled_table(doc, ["Технология", "Назначение"], frontend, col_widths=[7, 9])

    add_heading(doc, "4.3. База данных и инфраструктура", level=2)
    infra = [
        ("PostgreSQL 16 Alpine", "Реляционная СУБД"),
        ("Docker & Docker Compose", "Контейнеризация и развёртывание"),
        ("Git", "Система контроля версий"),
    ]
    add_styled_table(doc, ["Технология", "Назначение"], infra, col_widths=[7, 9])

    # ========== 5. МОДУЛЬНАЯ СТРУКТУРА ==========
    doc.add_page_break()
    add_heading(doc, "5. Модульная структура PMMIS", level=1)
    doc.add_paragraph(
        "Система состоит из 14 функциональных модулей, организованных в три категории:"
    )

    modules_core = [
        ("Dashboard", "Панель управления и KPI"),
        ("Контракты", "Реестр контрактов"),
        ("Подрядчики", "Реестр подрядчиков"),
        ("Проекты", "Структура проектов и компонентов"),
        ("План закупок", "Управление закупками"),
        ("Платежи", "Управление платежами"),
        ("Прогресс работ", "Физический прогресс по контрактам"),
        ("АВР", "Реестр актов выполненных работ"),
    ]
    modules_mgmt = [
        ("Задачи", "Управление задачами (Kanban/Calendar)"),
        ("Уведомления", "Многоканальные уведомления"),
    ]
    modules_ref = [
        ("География", "Районы, джамоаты, сёла"),
        ("Индикаторы / KPI", "Индикаторы проекта"),
        ("Справочники", "Категории и типы"),
        ("Пользователи / Роли", "Администрирование доступа"),
    ]

    add_heading(doc, "ОСНОВНОЕ (Core) — 8 модулей", level=3)
    add_styled_table(doc, ["Модуль", "Назначение"], modules_core, col_widths=[5, 11])

    add_heading(doc, "УПРАВЛЕНИЕ (Management) — 2 модуля", level=3)
    add_styled_table(doc, ["Модуль", "Назначение"], modules_mgmt, col_widths=[5, 11])

    add_heading(doc, "СПРАВОЧНИКИ (Reference Data) — 4 модуля", level=3)
    add_styled_table(doc, ["Модуль", "Назначение"], modules_ref, col_widths=[5, 11])

    add_heading(doc, "5.1. Страница авторизации", level=2)
    doc.add_paragraph(
        "При входе в систему пользователь попадает на страницу авторизации с "
        "профессиональным оформлением, содержащим логотип PMMIS и наименование проекта WSIP."
    )
    add_screenshot(doc, "login", "1 — Страница авторизации PMMIS")

    # ========== 6. ОПИСАНИЕ МОДУЛЕЙ ==========
    doc.add_page_break()
    add_heading(doc, "6. Описание модулей", level=1)

    # --- 6.1. Dashboard ---
    add_heading(doc, "6.1. Dashboard (Панель управления)", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("DashboardController")

    doc.add_paragraph(
        "Dashboard — это центральная панель управления, обеспечивающая мгновенный обзор "
        "состояния проекта. При входе пользователь видит агрегированные показатели и критические сигналы."
    )

    add_heading(doc, "KPI-карточки верхнего уровня:", level=3)
    kpi_top = [
        ("Общий бюджет", "Суммарный бюджет проекта (напр. $25 165 740)"),
        ("Сумма контрактов", "Суммарная стоимость всех контрактов"),
        ("Выплачено", "Общая сумма оплаченных платежей с процентом освоения"),
        ("Средний прогресс", "Средний процент физического выполнения по контрактам"),
    ]
    add_styled_table(doc, ["Карточка", "Описание"], kpi_top, col_widths=[5, 11])

    add_heading(doc, "Управленческие сигналы (Красные флаги):", level=3)
    doc.add_paragraph(
        "Система автоматически вычисляет и отображает критические управленческие сигналы на Dashboard:"
    )
    signals = [
        ("🔴 Просрочка контрактов", "Контракты с просроченными этапами (с указанием дней просрочки)"),
        ("🔴 Просроченные задачи", "Проверка АВР, подготовка документов и т.д."),
        ("⚠️ Задержки платежей", "Платежи, ожидающие одобрение длительное время"),
        ("⚠️ Задержки закупок", "Закупочные процедуры с отставанием от графика"),
    ]
    add_styled_table(doc, ["Сигнал", "Описание"], signals, col_widths=[5, 11])

    add_screenshot(doc, "dashboard", "2 — Dashboard PMMIS: KPI и управленческие сигналы")

    # --- 6.2. Проекты ---
    doc.add_page_break()
    add_heading(doc, "6.2. Модуль «Проекты»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("ProjectsController")

    doc.add_paragraph(
        "Модуль обеспечивает иерархическое управление структурой проектов по стандартам "
        "Всемирного Банка. Иерархия: Проект → Компонент → Подкомпонент."
    )

    add_heading(doc, "Функциональность:", level=3)
    funcs = [
        ("Создание и редактирование проектов", "Код проекта, название, бюджет, сроки"),
        ("Управление компонентами", "Каждый компонент имеет свой бюджет и описание"),
        ("Привязка контрактов", "Контракты привязываются к компонентам проекта"),
        ("Автоматический расчёт", "Общий бюджет агрегируется из компонентов"),
        ("Мультиязычность", "Названия на русском, таджикском и английском"),
    ]
    add_styled_table(doc, ["Функция", "Описание"], funcs, col_widths=[6, 10])

    doc.add_paragraph(
        "На текущий момент в системе зарегистрирован проект WSIP-1 с бюджетом $25 165 740, "
        "состоящий из 3 компонентов."
    )

    add_screenshot(doc, "projects", "3 — Модуль «Проекты»: структура проекта WSIP-1")

    # --- 6.3. Контракты ---
    doc.add_page_break()
    add_heading(doc, "6.3. Модуль «Контракты»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("ContractsController")

    doc.add_paragraph(
        "Модуль «Контракты» — это ядро системы PMMIS, вокруг которого строится вся "
        "управленческая цепочка. Реестр контрактов обеспечивает полную прозрачность исполнения."
    )

    add_heading(doc, "Основные данные контракта:", level=3)
    contract_fields = [
        ("Номер контракта", "Уникальный идентификатор (напр. RWSSP-W/006-02 Lot2)"),
        ("Объём работ", "Описание работ по контракту"),
        ("Подрядчик", "Привязка к реестру подрядчиков"),
        ("Тип контракта", "Работы / Консалтинг"),
        ("Дата подписания", "Дата подписания контракта"),
        ("Срок завершения", "Плановый срок завершения"),
        ("Сумма ($)", "Стоимость контракта в USD"),
        ("Прогресс (%)", "Процент физического выполнения"),
        ("Оплачено (%)", "Процент оплаченной суммы"),
    ]
    add_styled_table(doc, ["Поле", "Описание"], contract_fields, col_widths=[5, 11])

    add_heading(doc, "Дополнительные возможности:", level=3)
    extras = [
        ("Этапы контракта (Milestones)", "Контроль промежуточных дедлайнов"),
        ("Дополнения (Amendments)", "Фиксация изменений суммы или сроков"),
        ("Привязка индикаторов", "Contract-linked индикаторы для автоматического расчёта прогресса"),
        ("Мониторинг", "Визуальный мониторинг состояния всех контрактов"),
        ("Фильтрация", "По типу контракта, статусу, подрядчику, текстовый поиск"),
        ("Экспорт", "Excel Export и PDF Export"),
    ]
    add_styled_table(doc, ["Функция", "Описание"], extras, col_widths=[6, 10])

    add_screenshot(doc, "contracts", "4 — Модуль «Контракты»: реестр с фильтрами и экспортом")

    # --- 6.4. Подрядчики ---
    doc.add_page_break()
    add_heading(doc, "6.4. Модуль «Подрядчики»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("ContractorsController")

    doc.add_paragraph(
        "Централизованный реестр всех подрядчиков, участвующих в проекте."
    )

    contractor_fields = [
        ("Название", "Полное наименование организации"),
        ("Страна", "Страна регистрации"),
        ("Контактное лицо", "ФИО ответственного лица"),
        ("Email", "Электронная почта"),
        ("Телефон", "Контактный телефон"),
        ("Контрактов", "Количество привязанных контрактов (badge)"),
    ]
    add_styled_table(doc, ["Поле", "Описание"], contractor_fields, col_widths=[5, 11])

    doc.add_paragraph("В системе зарегистрировано 6 подрядчиков из Республики Таджикистан и Южной Кореи.")

    add_screenshot(doc, "contractors", "5 — Модуль «Подрядчики»: реестр с поиском и экспортом")

    # --- 6.5. АВР ---
    doc.add_page_break()
    add_heading(doc, "6.5. Модуль «АВР» (Акты выполненных работ)", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("WorkProgressReportsController")

    doc.add_paragraph(
        "Реестр АВР — это централизованное хранилище всех актов выполненных работ "
        "по всем контрактам. Модуль обеспечивает полную видимость хода работ."
    )

    add_heading(doc, "KPI-панель:", level=3)
    avr_kpi = [
        ("Всего АВР", "Общее количество зарегистрированных актов"),
        ("Контрактов", "По скольким контрактам есть АВР"),
        ("Средний прогресс", "Средний процент выполнения"),
    ]
    add_styled_table(doc, ["Показатель", "Описание"], avr_kpi, col_widths=[5, 11])

    add_heading(doc, "Автоматизация:", level=3)
    doc.add_paragraph("При создании АВР система автоматически:")
    auto_actions = [
        "Обновляет процент физического выполнения контракта",
        "Создаёт задачу «Проверить АВР по контракту [номер]» для ответственного сотрудника",
        "Отправляет уведомление куратору контракта",
    ]
    for action in auto_actions:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(action)

    add_screenshot(doc, "avr", "6 — Реестр АВР: KPI-панель, таблица с прогрессом")

    # --- 6.6. Прогресс работ ---
    doc.add_page_break()
    add_heading(doc, "6.6. Модуль «Прогресс работ»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("WorkProgressController")

    doc.add_paragraph(
        "Модуль предоставляет визуальную карточную раскладку прогресса работ по контрактам. "
        "В отличие от табличного реестра АВР, данный модуль отображает данные в формате "
        "карточек (Bootstrap Cards) с визуальными прогресс-барами."
    )

    add_heading(doc, "Workflow (рабочий процесс):", level=3)
    workflow = [
        ("Черновик", "Начальный статус, отчёт создан, но не отправлен"),
        ("На проверке", "Отправлен куратору для верификации"),
        ("Утверждён", "Проверен и подтверждён"),
    ]
    add_styled_table(doc, ["Статус", "Описание"], workflow, col_widths=[4, 12])

    add_screenshot(doc, "work_progress", "7 — Прогресс работ: карточный вид с workflow")

    # --- 6.7. Платежи ---
    doc.add_page_break()
    add_heading(doc, "6.7. Модуль «Платежи»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("PaymentsController")

    doc.add_paragraph(
        "Модуль управления платежами реализует полный workflow одобрения платежей "
        "в соответствии с требованиями Всемирного Банка."
    )

    add_heading(doc, "Финансовая панель:", level=3)
    fin_panel = [
        ("Всего", "Общая сумма зарегистрированных платежей"),
        ("Оплачено", "Сумма фактически оплаченных платежей"),
        ("Ожидает", "Сумма платежей, ожидающих одобрения (выделено красным)"),
    ]
    add_styled_table(doc, ["Показатель", "Описание"], fin_panel, col_widths=[4, 12])

    add_heading(doc, "Типы платежей:", level=3)
    pay_types = [
        ("Аванс", "Предоплата по контракту"),
        ("Промежуточный", "Оплата по АВР"),
        ("Финальный", "Финальный расчёт"),
        ("Удержание", "Гарантийное удержание"),
    ]
    add_styled_table(doc, ["Тип", "Описание"], pay_types, col_widths=[4, 12])

    add_heading(doc, "Workflow одобрения:", level=3)
    pay_wf = [
        ("Ожидает → Одобрен → Оплачен", "Стандартный путь платежа"),
        ("Ожидает → Отклонён", "С указанием причины отклонения"),
    ]
    add_styled_table(doc, ["Переход", "Описание"], pay_wf, col_widths=[6, 10])

    add_heading(doc, "Бизнес-валидация:", level=3)
    validations = [
        ("Hard Block", "Запрет оплаты сверх бюджета контракта"),
        ("Предупреждение 90%", "Сигнал при достижении 90% суммы контракта"),
        ("Missing AVR", "Предупреждение при отсутствии АВР для промежуточного платежа"),
        ("Аудит", "Автоматическая фиксация кто и когда одобрил/отклонил платёж"),
    ]
    add_styled_table(doc, ["Правило", "Описание"], validations, col_widths=[5, 11])

    add_screenshot(doc, "payments", "8 — Модуль «Платежи»: панель, workflow и фильтрация")

    # --- 6.8. План закупок ---
    doc.add_page_break()
    add_heading(doc, "6.8. Модуль «План закупок»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("ProcurementController")

    doc.add_paragraph("Модуль управления планом закупок по стандартам Всемирного Банка.")

    proc_fields = [
        ("Ref No", "Уникальный номер позиции"),
        ("Описание", "Описание закупки"),
        ("Проект", "Привязка к проекту"),
        ("Метод", "Метод закупки (ICB, NCB, Shopping и др.)"),
        ("Тип", "Тип закупки"),
        ("Сумма (USD)", "Плановая стоимость"),
        ("Статус", "Текущий статус процедуры"),
        ("Плановый тендер", "Плановая дата тендера"),
    ]
    add_styled_table(doc, ["Поле", "Описание"], proc_fields, col_widths=[5, 11])

    add_screenshot(doc, "procurement", "9 — Модуль «План закупок»: реестр с фильтрацией")

    # --- 6.9. Индикаторы ---
    doc.add_page_break()
    add_heading(doc, "6.9. Модуль «Индикаторы / KPI»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("IndicatorsController")

    doc.add_paragraph(
        "Модуль индикаторов реализован в точном соответствии с ТЗ и разделён на два класса индикаторов."
    )

    add_heading(doc, "A. Contract-linked индикаторы (автоматизируемые):", level=3)
    doc.add_paragraph(
        "Технические индикаторы, привязанные к контрактам. При создании АВР пользователь "
        "указывает прогресс по каждому привязанному индикатору. Система автоматически "
        "обновляет достигнутое значение и рассчитывает процент выполнения."
    )
    cl_indicators = [
        ("Кол-во установленных водомеров", "Число", "Привязан к контрактам"),
        ("Кол-во построенных РЧВ / водозаборов", "Число", "Привязан к контрактам"),
        ("Проложенные км трубопровода", "Число", "Привязан к контрактам"),
        ("Кол-во восстановленных скважин", "Число", "Привязан к контрактам"),
    ]
    add_styled_table(doc, ["Индикатор", "Тип", "Механизм"], cl_indicators, col_widths=[7, 2.5, 6.5])

    add_heading(doc, "B. Outcome-level индикаторы (ручной ввод):", level=3)
    doc.add_paragraph(
        "Индикаторы результатов WB (PDO, intermediate outcomes). Вводятся вручную "
        "через интерфейс. Система хранит историю и позволяет отслеживать динамику."
    )

    doc.add_paragraph("В системе зарегистрировано 18 индикаторов с иерархической структурой.")

    add_screenshot(doc, "indicators", "10 — Модуль «Индикаторы / KPI»: иерархический список")

    # --- 6.10. Задачи ---
    doc.add_page_break()
    add_heading(doc, "6.10. Модуль «Задачи»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("TasksController")

    doc.add_paragraph(
        "Модуль управления задачами реализован по образцу профессиональных систем "
        "управления проектами и представляет собой мультивидовой хаб с четырьмя перспективами."
    )

    add_heading(doc, "Перспективы (Views):", level=3)
    views = [
        ("📋 Список", "Табличный вид для массовых операций и высокоплотных данных"),
        ("📌 Доска", "Kanban-доска с drag-and-drop переключением статусов"),
        ("⏰ Сроки", "Фокус на дедлайнах и приоритетах"),
        ("📅 Календарь", "Месячный/недельный/дневной календарь (Syncfusion Scheduler)"),
    ]
    add_styled_table(doc, ["Вид", "Описание"], views, col_widths=[4, 12])

    add_heading(doc, "Автоматическое создание задач:", level=3)
    auto_tasks = [
        ("При создании АВР", "Задача «Проверить АВР по контракту [номер]»"),
        ("При создании платежа", "Задача «Подготовить документы к оплате #[номер]»"),
        ("При просрочке этапа контракта", "Задача с критическим приоритетом"),
    ]
    add_styled_table(doc, ["Событие", "Создаваемая задача"], auto_tasks, col_widths=[6, 10])

    add_screenshot(doc, "tasks", "11 — Модуль «Задачи»: календарный вид")

    # --- 6.11. Уведомления ---
    doc.add_page_break()
    add_heading(doc, "6.11. Модуль «Уведомления»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("NotificationsController")

    doc.add_paragraph(
        "Многоканальная система уведомлений, обеспечивающая информирование "
        "пользователей о критических событиях."
    )

    channels = [
        ("InApp", "Уведомления внутри системы с прямыми ссылками на объекты"),
        ("Email", "Отправка HTML-писем для формальных коммуникаций"),
        ("Telegram", "Интеграция для мобильных уведомлений"),
    ]
    add_styled_table(doc, ["Канал", "Описание"], channels, col_widths=[4, 12])

    add_heading(doc, "Фоновые службы:", level=3)
    services = [
        ("DeadlineNotificationService", "Автоматическая проверка приближающихся дедлайнов"),
        ("NotificationQueueService", "Очередь доставки уведомлений"),
    ]
    add_styled_table(doc, ["Служба", "Описание"], services, col_widths=[6, 10])

    # --- 6.12. География ---
    add_heading(doc, "6.12. Модуль «География»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллер: ")
    run.bold = True
    p.add_run("GeographyController")

    doc.add_paragraph(
        "Иерархическое управление географическими локациями проекта. "
        "Трёхуровневая структура: Район → Джамоат → Село/Кишлак."
    )

    geo_stats = [
        ("Районы", "3", "Балхинский, Вахшский, Дустийский"),
        ("Джамоаты", "13", "По 4-5 джамоата на район"),
        ("Населённые пункты", "68", "С данными о населении"),
    ]
    add_styled_table(doc, ["Уровень", "Количество", "Описание"], geo_stats, col_widths=[4, 3, 9])

    add_heading(doc, "Дополнительные возможности:", level=3)
    geo_features = [
        ("TreeView", "Альтернативный древовидный просмотр всей иерархии"),
        ("Drill-down навигация", "Переход от района к джамоату и далее к кишлаку"),
        ("Объекты инфраструктуры", "Школы и медучреждения с WSS-мониторингом"),
    ]
    add_styled_table(doc, ["Функция", "Описание"], geo_features, col_widths=[5, 11])

    add_screenshot(doc, "geography", "12 — Модуль «География»: иерархические справочники")

    # --- 6.13. Пользователи и Роли ---
    doc.add_page_break()
    add_heading(doc, "6.13. Модуль «Пользователи и Роли»", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Контроллеры: ")
    run.bold = True
    p.add_run("UsersController, RolesController, AccountController")

    doc.add_paragraph(
        "Система управления доступом на базе ASP.NET Core Identity "
        "со гранулярной матрицей разрешений."
    )

    roles = [
        ("PMU_ADMIN", "Полный доступ, управление пользователями и ролями"),
        ("PMU_STAFF", "Работа с данными проектов, контрактов, платежей"),
        ("WORLD_BANK", "Мониторинг и просмотр, аналитика"),
        ("CONTRACTOR", "Ограниченный доступ к своим контрактам"),
    ]
    add_styled_table(doc, ["Роль", "Уровень доступа"], roles, col_widths=[4, 12])

    add_heading(doc, "Матрица разрешений (RoleMenuPermission):", level=3)
    perms = [
        ("CanView", "Просмотр данных модуля"),
        ("CanCreate", "Создание новых записей"),
        ("CanEdit", "Редактирование существующих записей"),
        ("CanDelete", "Удаление записей"),
        ("CanViewAll", "Просмотр всех записей (не только своих)"),
        ("CanApprove", "Одобрение (для платежей)"),
    ]
    add_styled_table(doc, ["Разрешение", "Описание"], perms, col_widths=[4, 12])

    # --- 6.14. Справочники ---
    add_heading(doc, "6.14. Модуль «Справочники»", level=2)
    doc.add_paragraph(
        "Управление справочными данными системы: категории индикаторов (PDO, промежуточные, "
        "contract-linked), типы учреждений, типы оборудования, курсы валют."
    )

    # ========== 7. ВЗАИМОДЕЙСТВИЕ МОДУЛЕЙ ==========
    doc.add_page_break()
    add_heading(doc, "7. Взаимодействие модулей", level=1)

    doc.add_paragraph(
        "Модули PMMIS тесно интегрированы между собой, образуя единую управленческую цепочку."
    )

    add_heading(doc, "7.1. Ключевые связи между модулями", level=2)
    links = [
        ("Проекты → Закупки", "Валидация бюджета компонента", "Автоматическая"),
        ("Закупки → Контракты", "Создание контракта из завершённой закупки", "Ручная"),
        ("Контракты → АВР", "Привязка АВР к контракту", "Ручная"),
        ("АВР → Контракты", "Обновление % прогресса контракта", "Автоматическая"),
        ("АВР → Индикаторы", "Обновление contract-linked индикаторов", "Автоматическая"),
        ("АВР → Задачи", "Создание задачи «Проверить АВР»", "Автоматическая"),
        ("АВР → Уведомления", "Уведомление куратора об АВР", "Автоматическая"),
        ("Контракты → Платежи", "Привязка платежа к контракту", "Ручная"),
        ("Платежи → Задачи", "Создание задачи «Подготовить документы»", "Автоматическая"),
        ("Платежи → Уведомления", "Уведомление об одобрении/отклонении", "Автоматическая"),
        ("Контракты → Dashboard", "Агрегация KPI", "Автоматическая"),
        ("Все модули → Dashboard", "Управленческие сигналы (красные флаги)", "Автоматическая"),
    ]
    add_styled_table(doc, ["Из → В", "Тип связи", "Активация"], links, col_widths=[5, 7.5, 3.5])

    add_heading(doc, "7.2. Автоматические процессы", level=2)
    doc.add_paragraph("Система выполняет следующие автоматические действия без участия пользователя:")
    auto_procs = [
        ("1", "Обновление прогресса контракта при создании/изменении АВР"),
        ("2", "Обновление значений индикаторов при указании прогресса в АВР"),
        ("3", "Создание задач при создании АВР и платежей"),
        ("4", "Отправка уведомлений при смене статусов"),
        ("5", "Расчёт управленческих сигналов при каждом обращении к Dashboard"),
        ("6", "Проверка дедлайнов через DeadlineNotificationService"),
        ("7", "Доставка уведомлений через NotificationQueueService"),
    ]
    add_styled_table(doc, ["№", "Процесс"], auto_procs, col_widths=[1.5, 14.5])

    # ========== 8. СООТВЕТСТВИЕ ТЗ ==========
    doc.add_page_break()
    add_heading(doc, "8. Соответствие ТЗ", level=1)

    add_heading(doc, "8.1. Матрица соответствия требований", level=2)
    compliance = [
        ("3.1.1", "Контракты — реестр, сроки, стоимость, статус", "✅ Реализовано", "Контракты"),
        ("3.1.2", "АВР — загрузка файлов, статус проверки, связь с контрактом", "✅ Реализовано", "АВР, Прогресс работ"),
        ("3.1.3", "Платежи — запрос на оплату, статус выплаты, контроль обязательств", "✅ Реализовано", "Платежи"),
        ("3.1.4", "Обязательства и остатки — анализ расходов", "✅ Реализовано", "Контракты, Dashboard"),
        ("3.1.5", "Закупки — план, стадия, задержки", "✅ Реализовано", "План закупок"),
        ("4A", "Contract-linked индикаторы — автоматический расчёт", "✅ Реализовано", "Индикаторы"),
        ("4B", "Outcome-level индикаторы — ручной ввод", "✅ Реализовано", "Индикаторы"),
        ("6", "KPI для управления проектами", "✅ Реализовано", "Dashboard"),
        ("7", "Архитектура — веб-клиент, SQL БД", "✅ Реализовано", "Вся система"),
        ("—", "Загрузка/выгрузка Excel/CSV", "✅ Реализовано", "Контракты, АВР"),
        ("—", "Бессрочное хранение данных", "✅ Реализовано", "PostgreSQL"),
        ("8.1", "Не ERP — без бухгалтерии и кадров", "✅ Соблюдено", "—"),
        ("8.3", "Индикаторы двух типов", "✅ Соблюдено", "Индикаторы"),
        ("8.5", "Минимальная нагрузка на пользователей", "✅ Реализовано", "Авто-задачи"),
    ]
    add_styled_table(doc, ["Пункт ТЗ", "Требование", "Статус", "Модуль"], compliance, col_widths=[2, 7, 3.5, 3.5])

    add_heading(doc, "8.2. Дополнительно реализованные возможности", level=2)
    doc.add_paragraph("Помимо требований ТЗ, дополнительно реализованы:")
    additional = [
        ("Многоязычность", "Полная поддержка 3 языков (рус/тадж/англ)"),
        ("Модуль задач", "Полноценный Task Management с Kanban и Calendar"),
        ("Многоканальные уведомления", "InApp + Email + Telegram"),
        ("Управленческие сигналы", "Автоматические красные флаги на Dashboard"),
        ("Матрица разрешений", "Гранулярный контроль доступа по ролям"),
        ("Модуль географии", "Иерархические справочники районов/джамоатов/кишлаков"),
        ("Объекты инфраструктуры", "Школы и медучреждения с WSS-мониторингом"),
    ]
    add_styled_table(doc, ["Функциональность", "Описание"], additional, col_widths=[6, 10])

    # ========== 9. ЗАКЛЮЧЕНИЕ ==========
    doc.add_page_break()
    add_heading(doc, "9. Заключение", level=1)

    doc.add_paragraph(
        "В результате выполненных работ разработана и введена в эксплуатацию информационная "
        "система PMMIS, полностью соответствующая Финальной Концепции Системы и требованиям "
        "Всемирного Банка."
    )

    add_heading(doc, "Ключевые результаты:", level=2)

    results = [
        ("14 функциональных модулей", "Полный охват управленческой цепочки"),
        ("Ядро системы", "Прозрачная цепочка «Контракт → АВР → Платёж → Обязательства → Остаток»"),
        ("Автоматическая аналитика", "KPI, управленческие сигналы, прогнозирование рисков"),
        ("Профессиональный интерфейс", "Enterprise-grade UI на базе Syncfusion EJ2"),
        ("Многоязычность", "Полная локализация на 3 языка"),
        ("Масштабируемая архитектура", "Clean Architecture, Docker, PostgreSQL"),
        ("Минимальная нагрузка на пользователей", "Автоматические задачи, уведомления и расчёт показателей"),
    ]
    add_styled_table(doc, ["Результат", "Описание"], results, col_widths=[6, 10])

    doc.add_paragraph()
    conclusion = doc.add_paragraph()
    conclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conclusion.add_run(
        "Система обеспечивает дисциплину, прозрачность и снижение рисков\n"
        "при реализации проектов водоснабжения, являясь реально используемым\n"
        "ежедневным инструментом для PMU и ЦУП."
    )
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_paragraph()
    doc.add_paragraph()
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = signature.add_run("DonishSoft — Февраль 2026")
    run.bold = True
    run.font.size = Pt(12)

    # ========== SAVE ==========
    doc.save(OUTPUT_PATH)
    print(f"✅ Отчёт сохранён: {OUTPUT_PATH}")
    print(f"📊 Таблиц: ~30+")
    print(f"🖼️ Скриншотов: 12")


if __name__ == "__main__":
    create_report()
